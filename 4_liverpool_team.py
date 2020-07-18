import requests
import re
import bs4
import io
from urllib.request import urlopen
from PIL import Image
from docx import Document
from docx.shared import Inches

doc = Document()
header = doc.add_heading('Liverpool Team',0).add_run().bold = True
h1 = doc.add_table(rows = 1, cols = 3)
h1.style = 'Table Grid'
fields = h1.rows[0].cells
fields[0].text = 'name'
fields[1].text = 'link'
fields[2].text = 'image'

a = requests.get('https://www.liverpoolfc.com/team/first-team')
b = bs4.BeautifulSoup(a.text, 'html.parser')
c = b.select('li.team-player-list-item')

prefix = 'https://www.liverpoolfc.com'
d = []
for i in c:
    name = i.find('img')['alt']
    link = prefix + i.find('a')['href']
    image_from_url = urlopen(i.find('img')['src'])
    io_url = io.BytesIO()
    io_url.write(image_from_url.read())
    io_url.seek(0)

    fields = h1.add_row().cells
    fields[0].text = name
    fields[1].text = link
    image_field = fields[2].add_paragraph('').add_run().add_picture(io_url ,width = Inches(1.5))
    #image_field = fields[2].add_paragraph('').add_run().add_picture(img, width = Inches(0.5))
doc.save('liverpool members.docx')     