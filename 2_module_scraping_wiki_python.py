import requests
import re
import bs4
from docx import Document

def scrap_link(link,filename):
    doc = Document()
    doc.add_heading("{}".format(filename),0).add_run().bold = True
    h1 = doc.add_table(rows = 1, cols = 2)
    h1.style = 'Table Grid'
    fields = h1.rows[0].cells
    fields[0].text = 'Content'
    fields[1].text = 'Url Content'

    a = requests.get("{}".format(link))
    b = bs4.BeautifulSoup(a.text, 'html.parser')
    d = re.compile('^tocsection-')
    e = b.find_all('li', attrs={'class':d})
    prefix = "{}".format(link)

    content = []
    for j in e:
        link = prefix + j.find('a')['href']
        content.append(link)

    for g,i in zip(e,content):
        fields=h1.add_row().cells
        d=fields[0].add_paragraph(0).add_run(g.getText().split('\n')[0])
        if g['class'][0] == 'toclevel-1':
            d.bold =True
        fields[1].text = str(i)

    doc.save("{}.docx".format(filename))

if __name__=='__main__':
    scrap_link('https://en.wikipedia.org/wiki/Python_(programming_language)','Python Language')